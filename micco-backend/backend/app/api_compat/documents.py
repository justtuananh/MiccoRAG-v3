"""
Documents API — legacy compatibility layer for micco-server frontend.
Provides document CRUD, upload, versioning, thumbnails, and department-based access.
"""
from __future__ import annotations

import os
import uuid
import aiofiles
from pathlib import Path
from typing import Annotated

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy import select, func, or_, and_, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.deps import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.department import Department
from app.models.document import Document, DocumentStatus
from app.models.document_version import DocumentVersion
from docx import Document as DocxDocument
from app.schemas.compat import (
    LegacyDocumentVersionResponse,
    ProcessingStatusResponse,
    ProcessingStatusListResponse,
)
from app.api_compat.utils import (
    format_bytes_to_human,
    get_or_create_default_workspace,
    get_or_create_department_workspace,
    get_or_create_user_workspace,
    get_all_department_workspaces,
    map_rag_doc_to_legacy_with_dept,
    workspace_file_path,
)

ORG_APPROVER_ROLES = {"Admin", "Giám đốc", "Phó giám đốc"}
DEPT_APPROVER_ROLES = {"Admin", "Trưởng phòng"}

router = APIRouter(prefix="/api/documents", tags=["Documents"])

UPLOAD_DIR = settings.BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

THUMBNAIL_DIR = UPLOAD_DIR / "thumbnails"
THUMBNAIL_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".pptx", ".xlsx", ".png", ".jpg", ".jpeg"}


# ─── Access Control Helpers ────────────────────────────────────────

def _check_doc_access(user: User, doc: Document) -> None:
    """Raise 403 if user cannot access the document."""
    if user.role == "Admin":
        return
    # Public docs: tất cả user đều truy cập được
    if getattr(doc, "visibility", "internal") == "public":
        return
    # Uploader luôn truy cập được document của mình
    if getattr(doc, "uploader_id", None) == user.id:
        return
    # Private docs: chỉ uploader/Admin được truy cập
    if getattr(doc, "visibility", "internal") == "private":
        raise HTTPException(
            status_code=403,
            detail="Bạn không có quyền truy cập tài liệu cá nhân này",
        )
    # Private docs (personal workspace): chỉ uploader truy cập (đã check ở trên)
    # Department docs: phải cùng department
    if getattr(doc, "visibility", "internal") == "department":
        if doc.department_id is not None and user.department_id != doc.department_id:
            raise HTTPException(
                status_code=403,
                detail="Bạn không có quyền truy cập tài liệu này",
            )
        return
    # Internal: cùng department (legacy fallback)
    if doc.department_id is not None and user.department_id != doc.department_id:
        raise HTTPException(
            status_code=403,
            detail="Bạn không có quyền truy cập tài liệu này",
        )


def _can_modify_doc(user: User, doc: Document) -> None:
    """Raise 403 if user cannot modify/delete the document."""
    if user.role == "Admin":
        return
    if doc.uploader_id != user.id:
        raise HTTPException(
            status_code=403,
            detail="Chỉ người tải lên hoặc Admin mới có quyền thao tác với tài liệu này",
        )


# ─── Document Listing ───────────────────────────────────────────────

@router.get("")
async def list_documents(
    search: str | None = Query(None),
    type_filter: str | None = Query(None, alias="type"),
    category: str | None = Query(None),
    department_id: int | None = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List documents scoped to the user's department (admins see all)."""
    # NOTE: Documents are now stored per-department workspace (not just default workspace).
    # We list ALL documents (workspace filter removed) and rely on RBAC/department scoping below.

    # Base query with department + uploader joins
    stmt = (
        select(Document, Department.name.label("dept_name"), User.name.label("owner_name"))
        .outerjoin(Department, Document.department_id == Department.id)
        .outerjoin(User, Document.uploader_id == User.id)
    )

    # Department + visibility scoping
    if current_user.role != "Admin":
        stmt = stmt.where(
            or_(
                Document.visibility == "public",
                and_(
                    Document.visibility.in_(["internal", "department"]),
                    Document.department_id == current_user.department_id,
                ),
                Document.uploader_id == current_user.id,
            )
        )

    # Filter by selected department (admin only)
    if department_id is not None and current_user.role == "Admin":
        stmt = stmt.where(Document.department_id == department_id)

    # Filter by approval_status:
    # - Admin: sees approved + rejected (pending handled in approvals tab)
    # - Trưởng phòng: sees approved only
    # - Nhân viên: sees approved + their own pending, but never rejected
    if current_user.role == "Admin":
        stmt = stmt.where(Document.approval_status.notin_(["pending", "pending_org"]))
    elif current_user.role == "Trưởng phòng":
        stmt = stmt.where(Document.approval_status == "approved")
    else:
        stmt = stmt.where(
            or_(
                Document.approval_status == "approved",
                and_(
                    Document.uploader_id == current_user.id,
                    Document.approval_status.in_(["pending", "pending_org"]),
                ),
            )
        )

    # In-memory filters (for simplicity)
    rows = (await db.execute(stmt.order_by(Document.created_at.desc()))).all()
    mapped = []
    for row in rows:
        doc = row[0]
        dept_name = row[1]
        owner_name = row[2]
        mapped.append(map_rag_doc_to_legacy_with_dept(doc, owner_name=owner_name, dept_name=dept_name))

    # Apply search
    if search:
        s = search.lower()
        mapped = [d for d in mapped if s in d["name"].lower()]

    # Apply type filter
    if type_filter and type_filter != "All":
        mapped = [d for d in mapped if d["type"] == type_filter.upper()]

    # Apply category filter
    if category and category != "All":
        mapped = [d for d in mapped if d.get("category") == category]

    return mapped


# ─── Upload Documents ──────────────────────────────────────────────

@router.post("/upload")
async def upload_documents(
    files: list[UploadFile] = File(...),
    tags: str | None = Form(None),
    category: str | None = Form(None),
    visibility: str | None = Form("internal"),
    department_id: int | None = Form(None),
    thumbnail: UploadFile | None = File(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Upload one or more files.
    Admins/Trưởng phòng get auto-approved; regular users need approval.

    Workspace routing:
    - Tài liệu cá nhân (personal/private) → workspace cá nhân của uploader
    - Tài liệu nội bộ → workspace của phòng ban tác giả
    - Tài liệu công khai → workspace của phòng ban tác giả (sẽ được replicate khi approve)
    Nếu user không có phòng ban → fallback về default workspace.
    """
    requested_visibility = (visibility or "internal").lower()
    if requested_visibility in ("personal", "private"):
        effective_visibility = "private"
    elif requested_visibility in ("internal", "public", "department"):
        effective_visibility = requested_visibility
    else:
        effective_visibility = "internal"

    is_personal_upload = effective_visibility == "private"
    is_org_approver = current_user.role in ORG_APPROVER_ROLES
    is_dept_approver = current_user.role in DEPT_APPROVER_ROLES
    # Public docs require 2-level approval:
    # employee -> pending (dept), TP -> pending_org (org), Admin -> approved
    if is_personal_upload:
        doc_approval_status = "approved"
    elif effective_visibility == "public":
        if is_org_approver:
            doc_approval_status = "approved"
        elif is_dept_approver:
            doc_approval_status = "pending_org"
        else:
            doc_approval_status = "pending"
    else:
        doc_approval_status = "approved" if (is_org_approver or is_dept_approver) else "pending"
    effective_dept_id = department_id if department_id is not None else current_user.department_id

    # Resolve workspace: personal -> personal workspace, else department/default.
    if is_personal_upload:
        workspace = await get_or_create_user_workspace(db, current_user.id)
        effective_dept_id = None
    elif effective_dept_id:
        workspace = await get_or_create_department_workspace(db, effective_dept_id)
    else:
        workspace = await get_or_create_default_workspace(db)

    created: list[dict] = []

    for f in files:
        if f.size and f.size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File {f.filename} vượt quá giới hạn {MAX_FILE_SIZE // (1024*1024)}MB",
            )

        content = await f.read()
        ext = Path(f.filename or "file").suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Định dạng file không được hỗ trợ: {ext}",
            )

        # Save file
        stored_name = f"{uuid.uuid4().hex}{ext}"
        file_path = UPLOAD_DIR / stored_name
        async with aiofiles.open(file_path, "wb") as out:
            await out.write(content)

        # Handle thumbnail upload (shared across all files in batch)
        thumb_name = None
        if thumbnail:
            thumb_content = await thumbnail.read()
            thumb_ext = Path(thumbnail.filename or "file").suffix.lower()
            thumb_name = f"{uuid.uuid4().hex}{thumb_ext}"
            thumb_path = THUMBNAIL_DIR / thumb_name
            async with aiofiles.open(thumb_path, "wb") as out:
                await out.write(thumb_content)

        # Create document record
        doc = Document(
            workspace_id=workspace.id,
            filename=stored_name,
            original_filename=f.filename,
            file_type=ext[1:] if ext else "file",
            file_size=len(content),
            status=DocumentStatus.PENDING,
            uploader_id=current_user.id,
            department_id=effective_dept_id,
            visibility=effective_visibility,
            approval_status=doc_approval_status,
            category=category,
            tags=tags,
            thumbnail=thumb_name,
        )
        db.add(doc)
        await db.flush()

        # Create initial version V1.0
        version = DocumentVersion(
            document_id=doc.id,
            version_number=1,
            version_label="V 1.0",
            filename=stored_name,
            original_filename=f.filename,
            file_size=len(content),
            change_note="Phiên bản gốc",
            created_by=current_user.id,
            is_current=True,
        )
        db.add(version)

        # Auto-approved docs (admin/truongphong/personal): mark as PROCESSING and launch bg task
        if doc_approval_status == "approved":
            doc.status = DocumentStatus.PROCESSING
        await db.commit()

        # Launch background processing for auto-approved uploads
        if doc_approval_status == "approved":
            import asyncio
            from app.api.documents import process_document_background
            asyncio.get_event_loop().create_task(
                process_document_background(doc.id, str(file_path), workspace.id)
            )
            # If public: also index into all other department workspaces
            if effective_visibility == "public":
                other_workspaces = await get_all_department_workspaces(db)
                for other_ws in other_workspaces:
                    if other_ws.id != workspace.id:
                        asyncio.get_event_loop().create_task(
                            process_document_background(doc.id, str(file_path), other_ws.id)
                        )

        # Reload with department + uploader
        result = await db.execute(
            select(Document, Department.name.label("dept_name"), User.name.label("owner_name"))
            .outerjoin(Department, Document.department_id == Department.id)
            .outerjoin(User, Document.uploader_id == User.id)
            .where(Document.id == doc.id)
        )
        row = result.one_or_none()
        if row:
            doc, dept_name, owner_name = row
            created.append(map_rag_doc_to_legacy_with_dept(doc, owner_name=owner_name, dept_name=dept_name))
        else:
            created.append(map_rag_doc_to_legacy_with_dept(doc, owner_name=current_user.name))

    return created


# ─── Processing Status ────────────────────────────────────

@router.get("/processing-status", response_model=ProcessingStatusListResponse)
async def get_processing_status(
    filter: str = Query("all", description="Filter group: all | processing | indexed | failed"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Get documents by processing status group.
    - `all`        → pending / parsing / processing / indexing (active only)
    - `processing` → same as `all`
    - `indexed`    → completed (INDEXED)
    - `failed`     → failed (FAILED)

    Always returns `counts` with totals for all groups (for tab badges).

    NOTE: this route MUST be defined BEFORE /{doc_id} routes to prevent
    FastAPI from trying to parse "processing-status" as an integer.
    """
    from app.schemas.compat import StatusCounts
    from sqlalchemy import case

    # ── Determine which statuses to fetch ────────────────────────
    ACTIVE_STATUSES = ["PENDING", "PARSING", "PROCESSING", "INDEXING"]

    if filter in ("all", "processing"):
        fetch_statuses = ACTIVE_STATUSES
    elif filter == "indexed":
        fetch_statuses = ["INDEXED"]
    elif filter == "failed":
        fetch_statuses = ["FAILED"]
    else:
        fetch_statuses = ACTIVE_STATUSES  # fallback

    # ── Base access-scope sub-filter ─────────────────────────────
    if current_user.role in ("Admin", "Trưởng phòng"):
        if current_user.department_id:
            scope_filter = or_(
                Document.uploader_id == current_user.id,
                Document.department_id == current_user.department_id,
            )
        else:
            scope_filter = True  # Admin without dept → see all
    else:
        scope_filter = Document.uploader_id == current_user.id

    # ── Fetch items for requested filter ─────────────────────────
    stmt = (
        select(Document, Department.name.label("dept_name"), User.name.label("uploader_name"))
        .outerjoin(Department, Document.department_id == Department.id)
        .outerjoin(User, Document.uploader_id == User.id)
        .where(
            Document.status.in_(fetch_statuses),
            Document.approval_status != "rejected",
            scope_filter,
        )
        .order_by(Document.created_at.desc())
    )
    result = await db.execute(stmt)
    rows = result.all()

    items = [
        ProcessingStatusResponse(
            id=doc.id,
            name=doc.original_filename or doc.filename,
            status=doc.status.value if hasattr(doc.status, "value") else doc.status,
            chunk_count=doc.chunk_count or 0,
            error_message=doc.error_message,
            uploader_name=uploader_name or "Không rõ",
            department_name=dept_name,
            created_at=doc.created_at,
            file_type=doc.file_type,
            file_size=doc.file_size or 0,
        )
        for doc, dept_name, uploader_name in rows
    ]

    # ── Compute counts for ALL groups (for tab badges) ────────────
    count_stmt = (
        select(
            func.count(Document.id).filter(Document.status.in_(ACTIVE_STATUSES)).label("processing"),
            func.count(Document.id).filter(Document.status == "INDEXED").label("indexed"),
            func.count(Document.id).filter(Document.status == "FAILED").label("failed"),
        )
        .where(
            Document.approval_status != "rejected",
            scope_filter,
        )
    )
    count_row = (await db.execute(count_stmt)).one()
    cnt_processing = count_row.processing or 0
    cnt_indexed = count_row.indexed or 0
    cnt_failed = count_row.failed or 0

    counts = StatusCounts(
        all=cnt_processing,          # "Tất cả" badge = active docs
        processing=cnt_processing,
        indexed=cnt_indexed,
        failed=cnt_failed,
    )

    return ProcessingStatusListResponse(items=items, total=len(items), counts=counts)



# ─── Get Single Document ──────────────────────────────────────────

@router.get("/{doc_id}")
async def get_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get a single document by ID with department access check."""
    result = await db.execute(
        select(Document, Department.name.label("dept_name"), User.name.label("owner_name"))
        .outerjoin(Department, Document.department_id == Department.id)
        .outerjoin(User, Document.uploader_id == User.id)
        .where(Document.id == doc_id)
    )
    row = result.one_or_none()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    # Unpack carefully — owner_name may be None if uploader was deleted
    doc = row[0]
    dept_name = row[1]
    owner_name = row[2]

    _check_doc_access(current_user, doc)

    return map_rag_doc_to_legacy_with_dept(doc, owner_name=owner_name, dept_name=dept_name)


# ─── Download Document ────────────────────────────────────────────

@router.get("/{doc_id}/download")
async def download_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download the current version of a document."""
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _check_doc_access(current_user, doc)

    file_path = UPLOAD_DIR / doc.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    return FileResponse(
        path=str(file_path),
        filename=doc.original_filename or doc.filename,
        media_type="application/octet-stream",
    )


# ─── Preview Text Document ─────────────────────────────────────────

from docx import Document as DocxDocument

@router.get("/{doc_id}/preview")
async def preview_document_text(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Preview first few paragraphs/chars of text-based documents (docx, txt, md).
    Respects access controls.
    """
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _check_doc_access(current_user, doc)

    file_path = UPLOAD_DIR / doc.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = doc.file_type.lower()
    text_content = ""

    try:
        if ext == "docx":
            d = DocxDocument(str(file_path))
            text_content = "\n".join([p.text for p in d.paragraphs[:30]])
            if len(d.paragraphs) > 30:
                text_content += "\n\n...(Còn tiếp)..."
        elif ext in ("txt", "md"):
            async with aiofiles.open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                text_content = await f.read(5000)
                if len(text_content) == 5000:
                    text_content += "\n\n...(Còn tiếp)..."
        else:
            return {"supported": False, "message": "Preview not supported for this type"}

        return {"supported": True, "content": text_content, "file_type": ext}
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Error previewing file {file_path}: {str(e)}")
        return {"supported": False, "message": f"Lỗi rách tệp hoặc mất nội dung: {str(e)}"}


# ─── Document Thumbnail ────────────────────────────────────────────

@router.get("/{doc_id}/thumbnail")
async def get_thumbnail(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Serve document thumbnail image.
    Returns 404 if no thumbnail is associated with this document.
    """
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _check_doc_access(current_user, doc)

    # Thumbnail is stored in the document record as a path string
    thumb_path_str = getattr(doc, "thumbnail", None) or ""
    if not thumb_path_str:
        raise HTTPException(status_code=404, detail="Thumbnail not found")

    thumb_path = THUMBNAIL_DIR / thumb_path_str
    if not thumb_path.exists():
        raise HTTPException(status_code=404, detail="Thumbnail file not found on disk")

    ext = thumb_path_str.rsplit(".", 1)[-1].lower()
    media_types = {
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "gif": "image/gif",
        "webp": "image/webp",
    }
    media_type = media_types.get(ext, "image/jpeg")

    return FileResponse(path=str(thumb_path), media_type=media_type)


# ─── Document Versions ────────────────────────────────────────────

@router.get("/{doc_id}/versions", response_model=list[LegacyDocumentVersionResponse])
async def get_document_versions(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all versions of a document."""
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _check_doc_access(current_user, doc)

    versions = (
        await db.execute(
            select(DocumentVersion)
            .options(selectinload(DocumentVersion.creator))
            .where(DocumentVersion.document_id == doc_id)
            .order_by(DocumentVersion.version_number.desc())
        )
    ).scalars().all()

    return [
        LegacyDocumentVersionResponse(
            id=v.id,
            document_id=v.document_id,
            version_number=v.version_number,
            version_label=v.version_label,
            filename=v.filename,
            size=v.size_human,
            change_note=v.change_note,
            created_by_name=v.creator_name,
            is_current=bool(v.is_current),
            created_at=v.created_at,
        )
        for v in versions
    ]


@router.post("/{doc_id}/versions", response_model=LegacyDocumentVersionResponse)
async def upload_new_version(
    doc_id: int,
    file: UploadFile = File(...),
    change_note: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Upload a new version of an existing document.
    - Marks all previous versions as non-current.
    - Creates a new version record.
    - Updates the document's pointer to the latest file.
    """
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _can_modify_doc(current_user, doc)

    # Read file
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File vượt quá giới hạn {MAX_FILE_SIZE // (1024*1024)}MB",
        )

    ext = Path(file.filename or "file").suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Định dạng file không được hỗ trợ: {ext}",
        )

    # Save new file
    stored_name = f"{uuid.uuid4().hex}{ext}"
    file_path = UPLOAD_DIR / stored_name
    async with aiofiles.open(file_path, "wb") as out:
        await out.write(content)

    # Get next version number
    max_ver_row = await db.execute(
        select(func.max(DocumentVersion.version_number)).where(
            DocumentVersion.document_id == doc_id
        )
    )
    max_ver = max_ver_row.scalar() or 0
    next_ver = max_ver + 1

    # Mark old current versions as non-current
    await db.execute(
        update(DocumentVersion)
        .where(DocumentVersion.document_id == doc_id, DocumentVersion.is_current == True)
        .values(is_current=False)
    )

    # Create new version
    new_version = DocumentVersion(
        document_id=doc_id,
        version_number=next_ver,
        version_label=f"V {next_ver}.0",
        filename=stored_name,
        original_filename=file.filename,
        file_size=len(content),
        change_note=change_note or f"Phiên bản {next_ver}.0",
        created_by=current_user.id,
        is_current=True,
    )
    db.add(new_version)

    # Update document pointer to new file
    doc.filename = stored_name
    doc.original_filename = file.filename
    doc.file_size = len(content)
    doc.file_type = ext[1:] if ext else "file"
    is_org_approver = current_user.role in ORG_APPROVER_ROLES
    is_dept_approver = current_user.role in DEPT_APPROVER_ROLES
    is_personal_doc = (doc.visibility or "internal") == "private"
    if is_personal_doc:
        doc.approval_status = "approved"
    elif (doc.visibility or "internal") == "public":
        if is_org_approver:
            doc.approval_status = "approved"
        elif is_dept_approver:
            doc.approval_status = "pending_org"
        else:
            doc.approval_status = "pending"
    else:
        doc.approval_status = "approved" if (is_org_approver or is_dept_approver) else "pending"
    doc.status = DocumentStatus.PROCESSING if doc.approval_status == "approved" else DocumentStatus.PENDING

    await db.commit()
    await db.refresh(new_version)

    if doc.approval_status == "approved":
        import asyncio
        from app.api.documents import process_document_background

        asyncio.get_event_loop().create_task(
            process_document_background(doc.id, str(file_path), doc.workspace_id)
        )

    # Load creator
    result = await db.execute(
        select(DocumentVersion)
        .options(selectinload(DocumentVersion.creator))
        .where(DocumentVersion.id == new_version.id)
    )
    loaded = result.scalar_one()

    return LegacyDocumentVersionResponse(
        id=loaded.id,
        document_id=loaded.document_id,
        version_number=loaded.version_number,
        version_label=loaded.version_label,
        filename=loaded.filename,
        size=loaded.size_human,
        change_note=loaded.change_note,
        created_by_name=loaded.creator_name,
        is_current=bool(loaded.is_current),
        created_at=loaded.created_at,
    )


@router.get("/{doc_id}/versions/{version_id}/download")
async def download_version(
    doc_id: int,
    version_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download a specific version of a document."""
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _check_doc_access(current_user, doc)

    version = (
        await db.execute(
            select(DocumentVersion).where(
                DocumentVersion.id == version_id,
                DocumentVersion.document_id == doc_id,
            )
        )
    ).scalar_one_or_none()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    file_path = UPLOAD_DIR / version.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Version file not found on disk")

    return FileResponse(
        path=str(file_path),
        filename=version.original_filename or f"{doc.original_filename} ({version.version_label})",
        media_type="application/octet-stream",
    )


# ─── Delete Document ───────────────────────────────────────────────

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Delete a document and all its versions.
    Only the uploader or Admin can delete.
    """
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _can_modify_doc(current_user, doc)

    # Delete file from disk
    file_path = UPLOAD_DIR / doc.filename
    if file_path.exists():
        os.remove(file_path)

    # Delete all version files
    versions = (
        await db.execute(
            select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
        )
    ).scalars().all()
    for v in versions:
        vp = UPLOAD_DIR / v.filename
        if vp.exists():
            os.remove(vp)

    await db.delete(doc)
    await db.commit()
    return {"message": "Xóa tài liệu thành công"}

