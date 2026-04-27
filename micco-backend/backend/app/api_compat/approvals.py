from __future__ import annotations
from fastapi import APIRouter, Depends, HTTPException, Body, BackgroundTasks
from sqlalchemy import select, update, func, or_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.deps import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.document import Document, DocumentStatus
from app.models.knowledge_entry import KnowledgeEntry
from app.models.department import Department
from app.api_compat.utils import (
    format_bytes_to_human,
    get_or_create_department_workspace,
    get_all_department_workspaces,
    get_or_create_default_workspace,
)
from app.api.documents import process_document_background, process_knowledge_background, UPLOAD_DIR
from docx import Document as DocxDocument
import aiofiles
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/approvals", tags=["Approvals"])

ORG_APPROVER_ROLES = {"Admin", "Giám đốc", "Phó giám đốc"}
DEPT_APPROVER_ROLES = {"Admin", "Trưởng phòng"}
ALL_APPROVER_ROLES = ORG_APPROVER_ROLES | DEPT_APPROVER_ROLES

DOC_PENDING_DEPT = "pending"
DOC_PENDING_ORG = "pending_org"
DOC_APPROVED = "approved"

KN_PENDING_DEPT = "pending_dept"
KN_PENDING_ORG = "pending_org"
KN_PENDING_LEGACY = "pending_approval"
KN_APPROVED = "approved"
KN_REJECTED = "rejected"
KN_VISIBILITY_PUBLIC = "public"
KN_VISIBILITY_PRIVATE = "private"


@router.get("/count")
async def pending_count(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        return {"count": 0}

    if current_user.role == "Admin":
        doc_stmt = select(func.count(Document.id)).where(
            or_(
                Document.approval_status == DOC_PENDING_DEPT,
                Document.approval_status == DOC_PENDING_ORG,
            )
        )
    elif current_user.role == "Trưởng phòng":
        doc_stmt = select(func.count(Document.id)).where(
            Document.approval_status == DOC_PENDING_DEPT,
            Document.department_id == current_user.department_id,
        )
    else:
        # Ban giám đốc chỉ nhận duyệt cấp tổ chức cho tài liệu công khai
        doc_stmt = select(func.count(Document.id)).where(
            Document.approval_status == DOC_PENDING_ORG,
            Document.visibility == KN_VISIBILITY_PUBLIC,
        )
    doc_count = (await db.execute(doc_stmt)).scalar() or 0

    if current_user.role == "Admin":
        kn_stmt = select(func.count(KnowledgeEntry.id)).where(
            or_(
                KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
                KnowledgeEntry.approval_status == KN_PENDING_ORG,
            )
        )
    elif current_user.role == "Trưởng phòng":
        kn_stmt = select(func.count(KnowledgeEntry.id)).where(
            KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
            KnowledgeEntry.department_id == current_user.department_id,
        )
    else:
        kn_stmt = select(func.count(KnowledgeEntry.id)).where(
            KnowledgeEntry.approval_status == KN_PENDING_ORG,
            KnowledgeEntry.visibility == KN_VISIBILITY_PUBLIC,
        )
    kn_count = (await db.execute(kn_stmt)).scalar() or 0

    # Always return last_requester so frontend can show who uploaded
    # even when count increases from 0 → N
    last_requester = None

    # Fetch the most recent pending item (doc or knowledge) scoped by approval stage
    if current_user.role == "Admin":
        latest_doc_result = await db.execute(
            select(User.name, Document.created_at)
            .join(User, Document.uploader_id == User.id)
            .where(
                or_(
                    Document.approval_status == DOC_PENDING_DEPT,
                    Document.approval_status == DOC_PENDING_ORG,
                )
            )
            .order_by(Document.created_at.desc())
            .limit(1)
        )
    elif current_user.role == "Trưởng phòng":
        latest_doc_result = await db.execute(
            select(User.name, Document.created_at)
            .join(User, Document.uploader_id == User.id)
            .where(
                Document.approval_status == DOC_PENDING_DEPT,
                Document.department_id == current_user.department_id,
            )
            .order_by(Document.created_at.desc())
            .limit(1)
        )
    else:
        latest_doc_result = await db.execute(
            select(User.name, Document.created_at)
            .join(User, Document.uploader_id == User.id)
            .where(
                Document.approval_status == DOC_PENDING_ORG,
                Document.visibility == KN_VISIBILITY_PUBLIC,
            )
            .order_by(Document.created_at.desc())
            .limit(1)
        )
    doc_row = latest_doc_result.first()

    if current_user.role == "Admin":
        latest_kn_result = await db.execute(
            select(User.name, KnowledgeEntry.created_at)
            .join(User, KnowledgeEntry.owner_id == User.id)
            .where(
                or_(
                    KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
                    KnowledgeEntry.approval_status == KN_PENDING_ORG,
                )
            )
            .order_by(KnowledgeEntry.created_at.desc())
            .limit(1)
        )
    elif current_user.role == "Trưởng phòng":
        latest_kn_result = await db.execute(
            select(User.name, KnowledgeEntry.created_at)
            .join(User, KnowledgeEntry.owner_id == User.id)
            .where(
                KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
                KnowledgeEntry.department_id == current_user.department_id,
            )
            .order_by(KnowledgeEntry.created_at.desc())
            .limit(1)
        )
    else:
        latest_kn_result = await db.execute(
            select(User.name, KnowledgeEntry.created_at)
            .join(User, KnowledgeEntry.owner_id == User.id)
            .where(
                KnowledgeEntry.approval_status == KN_PENDING_ORG,
                KnowledgeEntry.visibility == KN_VISIBILITY_PUBLIC,
            )
            .order_by(KnowledgeEntry.created_at.desc())
            .limit(1)
        )
    kn_row = latest_kn_result.first()

    if doc_row and kn_row:
        last_requester = doc_row[0] if doc_row[1] > kn_row[1] else kn_row[0]
    elif doc_row:
        last_requester = doc_row[0]
    elif kn_row:
        last_requester = kn_row[0]

    total = doc_count + kn_count
    return {"count": total, "last_requester": last_requester}



@router.get("/pending")
async def list_pending(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        return {"documents": [], "knowledge": []}

    # Join with User and Department to get names
    if current_user.role == "Admin":
        stmt = (
            select(Document, User.name, Department.name)
            .outerjoin(User, Document.uploader_id == User.id)
            .outerjoin(Department, Document.department_id == Department.id)
            .where(
                or_(
                    Document.approval_status == DOC_PENDING_DEPT,
                    Document.approval_status == DOC_PENDING_ORG,
                )
            )
            .order_by(Document.created_at.desc())
        )
    elif current_user.role == "Trưởng phòng":
        stmt = (
            select(Document, User.name, Department.name)
            .outerjoin(User, Document.uploader_id == User.id)
            .outerjoin(Department, Document.department_id == Department.id)
            .where(
                Document.approval_status == DOC_PENDING_DEPT,
                Document.department_id == current_user.department_id,
            )
            .order_by(Document.created_at.desc())
        )
    else:
        stmt = (
            select(Document, User.name, Department.name)
            .outerjoin(User, Document.uploader_id == User.id)
            .outerjoin(Department, Document.department_id == Department.id)
            .where(
                Document.approval_status == DOC_PENDING_ORG,
                Document.visibility == KN_VISIBILITY_PUBLIC,
            )
            .order_by(Document.created_at.desc())
        )
    result = await db.execute(stmt)
    rows = result.all()

    docs = []
    for doc, uploader_name, dept_name in rows:
        docs.append({
            "id": doc.id,
            "name": doc.original_filename,
            "category": "Tài liệu",
            "owner": uploader_name or "Hệ thống",
            "department": dept_name or "Chung",
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "size": format_bytes_to_human(doc.file_size or 0),
            "visibility": doc.visibility,
            "file_type": doc.file_type.lower(),
            "approval_status": doc.approval_status,
        })

    # Fetch pending knowledge entries by approval stage
    if current_user.role == "Admin":
        kn_stmt = (
            select(KnowledgeEntry, User.name, Department.name)
            .outerjoin(User, KnowledgeEntry.owner_id == User.id)
            .outerjoin(Department, KnowledgeEntry.department_id == Department.id)
            .where(
                or_(
                    KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
                    KnowledgeEntry.approval_status == KN_PENDING_ORG,
                )
            )
            .order_by(KnowledgeEntry.created_at.desc())
        )
    elif current_user.role == "Trưởng phòng":
        kn_stmt = (
            select(KnowledgeEntry, User.name, Department.name)
            .outerjoin(User, KnowledgeEntry.owner_id == User.id)
            .outerjoin(Department, KnowledgeEntry.department_id == Department.id)
            .where(
                KnowledgeEntry.approval_status.in_([KN_PENDING_DEPT, KN_PENDING_LEGACY]),
                KnowledgeEntry.department_id == current_user.department_id,
            )
            .order_by(KnowledgeEntry.created_at.desc())
        )
    else:
        kn_stmt = (
            select(KnowledgeEntry, User.name, Department.name)
            .outerjoin(User, KnowledgeEntry.owner_id == User.id)
            .outerjoin(Department, KnowledgeEntry.department_id == Department.id)
            .where(
                KnowledgeEntry.approval_status == KN_PENDING_ORG,
                KnowledgeEntry.visibility == KN_VISIBILITY_PUBLIC,
            )
            .order_by(KnowledgeEntry.created_at.desc())
        )
    kn_result = await db.execute(kn_stmt)
    kn_rows = kn_result.all()

    knowledge_items = []
    for entry, owner_name, dept_name in kn_rows:
        knowledge_items.append({
            "id": entry.id,
            "title": entry.title,
            "content_text": entry.content_text,
            "category": entry.category,
            "owner": owner_name or "Hệ thống",
            "department": dept_name or "Chung",
            "created_at": entry.created_at.isoformat() if entry.created_at else None,
            "visibility": entry.visibility,
            "tags": entry.tags,
            "approval_status": entry.approval_status,
        })

    return {"documents": docs, "knowledge": knowledge_items}


@router.post("/documents/{doc_id}/approve")
async def approve_document(
    doc_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.approval_status == DOC_PENDING_DEPT:
        if current_user.role not in DEPT_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
            if doc.department_id != current_user.department_id:
                raise HTTPException(
                    status_code=403,
                    detail="Không thể phê duyệt tài liệu của phòng ban khác"
                )
        if (doc.visibility or "internal") == KN_VISIBILITY_PUBLIC:
            doc.approval_status = DOC_PENDING_ORG
            doc.status = DocumentStatus.PENDING
            await db.commit()
            return {
                "message": "Đã duyệt cấp phòng, đã thông báo tới Ban giám đốc",
                "id": doc_id,
                "processing_started": False,
            }
    elif doc.approval_status == DOC_PENDING_ORG:
        if current_user.role not in ORG_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if (doc.visibility or "internal") != KN_VISIBILITY_PUBLIC:
            raise HTTPException(status_code=400, detail="Tài liệu nội bộ chỉ cần Trưởng phòng phê duyệt")
    else:
        raise HTTPException(status_code=400, detail="Tài liệu không ở trạng thái chờ duyệt")

    doc.approval_status = DOC_APPROVED
    doc.status = DocumentStatus.PROCESSING
    await db.commit()

    # Resolve target workspace: use document's department workspace
    from app.core.config import settings as _settings
    file_path = str(_settings.BASE_DIR / "uploads" / doc.filename)

    if doc.department_id:
        target_ws = await get_or_create_department_workspace(db, doc.department_id)
    else:
        target_ws = await get_or_create_default_workspace(db)

    # Update document's workspace_id to match its department
    doc.workspace_id = target_ws.id
    await db.commit()

    # Trigger background parsing & indexing into department workspace
    background_tasks.add_task(process_document_background, doc.id, file_path, target_ws.id)

    # If public: replicate indexing into all OTHER department workspaces
    if doc.visibility == "public":
        other_workspaces = await get_all_department_workspaces(db)
        for other_ws in other_workspaces:
            if other_ws.id != target_ws.id:
                background_tasks.add_task(
                    process_document_background, doc.id, file_path, other_ws.id
                )

    return {"message": "Đã phê duyệt và đang bắt đầu xử lý", "id": doc_id, "processing_started": True}


@router.post("/documents/{doc_id}/reject")
async def reject_document(
    doc_id: int,
    note: str = Body(None, embed=True),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.approval_status == DOC_PENDING_DEPT:
        if current_user.role not in DEPT_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if current_user.role == "Trưởng phòng" and doc.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể từ chối tài liệu của phòng ban khác"
            )
    elif doc.approval_status == DOC_PENDING_ORG:
        if current_user.role not in ORG_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        # Ban giám đốc: chỉ từ chối ở bước cấp tổ chức của tài liệu công khai
        if (doc.visibility or "internal") != KN_VISIBILITY_PUBLIC:
            raise HTTPException(status_code=400, detail="Chỉ có thể từ chối tài liệu công khai ở bước duyệt cấp tổ chức")
    else:
        raise HTTPException(status_code=400, detail="Tài liệu không ở trạng thái chờ duyệt")

    doc.approval_status = "rejected"
    doc.approval_note = note
    doc.status = DocumentStatus.REJECTED
    await db.commit()
    return {"message": "Đã từ chối tài liệu", "id": doc_id}
@router.get("/documents/{doc_id}/status")
async def get_document_status(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Return current processing status of an approved document."""
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "id": doc.id,
        "status": doc.status.value if hasattr(doc.status, "value") else doc.status,
        "approval_status": doc.approval_status,
        "chunk_count": doc.chunk_count,
        "error_message": doc.error_message,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
    }


@router.get("/documents/{doc_id}/preview")
async def preview_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = UPLOAD_DIR / doc.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = doc.file_type.lower()
    text_content = ""

    try:
        if ext == "docx":
            # Quick sync read for docx (safe for small snippets)
            d = DocxDocument(str(file_path))
            # Get first 30 paragraphs
            text_content = "\n".join([p.text for p in d.paragraphs[:30]])
            if len(d.paragraphs) > 30:
                text_content += "\n\n...(Còn tiếp)..."
        elif ext in ("txt", "md"):
            async with aiofiles.open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                text_content = await f.read(5000) # first 5k chars
                if len(text_content) == 5000:
                    text_content += "\n\n...(Còn tiếp)..."
        else:
            return {"supported": False, "message": "Preview not supported for this type"}
            
        return {"supported": True, "content": text_content, "file_type": ext}
    except Exception as e:
        logger.error(f"Error previewing file {file_path}: {str(e)}")
        return {"supported": False, "message": f"Error loading preview: {str(e)}"}


@router.post("/knowledge/{entry_id}/approve")
async def approve_knowledge(
    entry_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    entry = (await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id)
    )).scalar_one_or_none()

    if not entry:
        raise HTTPException(status_code=404, detail="Knowledge entry not found")

    visibility = (entry.visibility or "internal").lower()
    if visibility in ("personal", KN_VISIBILITY_PRIVATE):
        raise HTTPException(status_code=400, detail="Tri thức cá nhân không cần phê duyệt")

    if entry.approval_status in (KN_PENDING_DEPT, KN_PENDING_LEGACY):
        if current_user.role not in DEPT_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if current_user.role == "Trưởng phòng" and entry.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể phê duyệt tri thức của phòng ban khác"
            )
        if visibility == KN_VISIBILITY_PUBLIC:
            entry.approval_status = KN_PENDING_ORG
            entry.status = "Pending"
            await db.commit()
            return {"message": "Đã duyệt cấp phòng, đang chờ duyệt cấp tổ chức", "id": entry_id}
    elif entry.approval_status == KN_PENDING_ORG:
        if current_user.role not in ORG_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if visibility != KN_VISIBILITY_PUBLIC:
            raise HTTPException(status_code=400, detail="Tri thức phòng ban cần Trưởng phòng phê duyệt")
    else:
        raise HTTPException(status_code=400, detail="Tri thức không ở trạng thái chờ duyệt")

    entry.approval_status = KN_APPROVED
    entry.status = "Active"
    entry.ingest_status = "processing"
    await db.commit()

    # Resolve workspace for this knowledge entry
    if entry.department_id:
        target_ws = await get_or_create_department_workspace(db, entry.department_id)
    else:
        target_ws = await get_or_create_default_workspace(db)

    # Trigger background indexing into the department workspace
    background_tasks.add_task(process_knowledge_background, entry_id, target_ws.id)

    # If public: replicate into all other department workspaces
    if getattr(entry, 'visibility', 'internal') == KN_VISIBILITY_PUBLIC:
        other_workspaces = await get_all_department_workspaces(db)
        for other_ws in other_workspaces:
            if other_ws.id != target_ws.id:
                background_tasks.add_task(process_knowledge_background, entry_id, other_ws.id)

    return {"message": "Đã phê duyệt tri thức", "id": entry_id}

@router.post("/knowledge/{entry_id}/reject")
async def reject_knowledge(
    entry_id: int,
    note: str = Body(None, embed=True),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ALL_APPROVER_ROLES:
        raise HTTPException(status_code=403, detail="Permission denied")

    entry = (await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id)
    )).scalar_one_or_none()

    if not entry:
        raise HTTPException(status_code=404, detail="Knowledge entry not found")

    rejection_note = (note or "").strip()
    if not rejection_note:
        raise HTTPException(status_code=400, detail="Vui lòng nhập lý do từ chối tri thức")

    visibility = (entry.visibility or "internal").lower()

    if entry.approval_status in (KN_PENDING_DEPT, KN_PENDING_LEGACY):
        if current_user.role not in DEPT_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if current_user.role == "Trưởng phòng" and entry.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể từ chối tri thức của phòng ban khác"
            )
    elif entry.approval_status == KN_PENDING_ORG:
        if current_user.role not in ORG_APPROVER_ROLES:
            raise HTTPException(status_code=403, detail="Permission denied")
        if visibility != KN_VISIBILITY_PUBLIC:
            raise HTTPException(status_code=400, detail="Chỉ có thể từ chối tri thức công khai ở bước duyệt cấp tổ chức")
    else:
        raise HTTPException(status_code=400, detail="Tri thức không ở trạng thái chờ duyệt")

    entry.approval_status = KN_REJECTED
    entry.approval_note = rejection_note
    entry.status = "Draft"
    await db.commit()
    return {"message": "Đã từ chối tri thức", "id": entry_id}
